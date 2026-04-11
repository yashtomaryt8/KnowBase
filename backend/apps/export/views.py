from docx import Document as DocxDocument
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from rest_framework.views import APIView

from apps.pages.models import Page
from apps.topics.models import Topic


def extract_text(node):
    text_parts = []

    if isinstance(node, dict):
        text_value = node.get('text')
        if text_value:
            text_parts.append(text_value)

        for child in node.get('content', []):
            child_text = extract_text(child)
            if child_text:
                text_parts.append(child_text)
    elif isinstance(node, list):
        for item in node:
            item_text = extract_text(item)
            if item_text:
                text_parts.append(item_text)

    return ''.join(text_parts)


def _apply_shading(paragraph, fill='F3F4F6'):
    paragraph_properties = paragraph._element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), fill)
    paragraph_properties.append(shading)


def tiptap_to_docx(doc, content_json):
    for node in content_json.get('content', []):
        node_type = node.get('type')

        if node_type == 'heading':
            level = node.get('attrs', {}).get('level', 1)
            level = min(max(int(level), 1), 9)
            doc.add_heading(extract_text(node), level=level)
        elif node_type == 'paragraph':
            text = extract_text(node).strip()
            if text:
                doc.add_paragraph(text)
        elif node_type == 'bulletList':
            for item in node.get('content', []):
                text = extract_text(item).strip()
                if text:
                    doc.add_paragraph(text, style='List Bullet')
        elif node_type == 'orderedList':
            for item in node.get('content', []):
                text = extract_text(item).strip()
                if text:
                    doc.add_paragraph(text, style='List Number')
        elif node_type == 'codeBlock':
            text = extract_text(node)
            if text:
                paragraph = doc.add_paragraph()
                _apply_shading(paragraph)
                run = paragraph.add_run(text)
                run.font.name = 'Courier New'
                run_properties = run._element.get_or_add_rPr()
                run_properties.rFonts.set(qn('w:eastAsia'), 'Courier New')
                run.font.size = Pt(9)
        elif node_type == 'blockquote':
            text = extract_text(node).strip()
            if text:
                doc.add_paragraph(text, style='Quote')


class ExportPageDocxView(APIView):
    def get(self, request, pk):
        page = get_object_or_404(Page.objects.select_related('topic'), pk=pk)
        doc = DocxDocument()

        doc.add_heading(page.title, level=0)
        metadata = doc.add_paragraph()
        metadata_run = metadata.add_run(
            f'Topic: {page.topic.name}  |  {page.created_at.strftime("%Y-%m-%d %H:%M")}'
        )
        metadata_run.font.color.rgb = RGBColor(128, 128, 128)

        tiptap_to_docx(doc, page.content_json)

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = (
            f'attachment; filename="{page.title.replace(" ", "_")}.docx"'
        )
        doc.save(response)
        return response


class ExportTopicDocxView(APIView):
    def get(self, request, pk):
        topic = get_object_or_404(Topic, pk=pk)
        pages = topic.pages.order_by('sort_order', 'created_at')
        doc = DocxDocument()

        for index, page in enumerate(pages):
            if index > 0:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

            doc.add_heading(page.title, level=1)
            tiptap_to_docx(doc, page.content_json)

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = (
            f'attachment; filename="{topic.name.replace(" ", "_")}.docx"'
        )
        doc.save(response)
        return response
